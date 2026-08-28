"""
Step 6: Render a DocumentPlan into a .docx file using python-docx.
No LLM involvement here — pure deterministic Python.
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from schemas import DocumentPlan


def _add_table(doc: Document, table_data) -> None:
    """Render a Table schema object as a real Word table."""
    if not table_data.headers:
        return
    col_count = len(table_data.headers)
    tbl = doc.add_table(rows=1, cols=col_count, style="Light Grid Accent 1")

    # Header row
    for i, header in enumerate(table_data.headers):
        cell = tbl.rows[0].cells[i]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Data rows
    for row_data in table_data.rows:
        row = tbl.add_row()
        for i, value in enumerate(row_data[:col_count]):
            row.cells[i].text = value

    doc.add_paragraph()  # spacer after table


def generate_docx(plan: DocumentPlan, output_path: str) -> str:
    doc = Document()

    # Title
    doc.add_heading(plan.title, level=0)

    # Subtitle
    if plan.subtitle:
        p = doc.add_paragraph(plan.subtitle)
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(13)

    # Sections
    for section in plan.sections:
        doc.add_heading(section.heading, level=1)

        for para_text in section.paragraphs:
            doc.add_paragraph(para_text)

        for bullet in section.bullets:
            doc.add_paragraph(bullet.text, style="List Bullet")
            for sub in bullet.sub_bullets:
                doc.add_paragraph(sub, style="List Bullet 2")

        for table_data in section.tables:
            _add_table(doc, table_data)

    doc.save(output_path)
    return output_path
